from __future__ import annotations

import io
import shutil
import textwrap
import time
import uuid
from pathlib import Path

import pymupdf
from docx import Document
from PIL import Image, ImageDraw, ImageFont, ImageOps

ALLOWED = {".pdf", ".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".txt", ".md", ".docx"}


class UnsupportedDocument(ValueError):
    pass


class DocumentStore:
    def __init__(self, root: Path):
        self.root = root
        self.root.mkdir(parents=True, exist_ok=True)

    def create(self, filename: str, content: bytes) -> tuple[str, list[Path]]:
        self._remove_expired()
        suffix = Path(filename).suffix.lower()
        if suffix not in ALLOWED:
            raise UnsupportedDocument(f"Unsupported file type: {suffix or 'unknown'}")
        doc_id = uuid.uuid4().hex
        target = self.root / doc_id
        target.mkdir()
        pages = self._render(suffix, content, target)
        return doc_id, pages

    def _remove_expired(self) -> None:
        cutoff = time.time() - 24 * 60 * 60
        for directory in self.root.iterdir():
            if directory.is_dir() and directory.stat().st_mtime < cutoff:
                shutil.rmtree(directory, ignore_errors=True)

    def pages(self, doc_id: str) -> list[Path]:
        if not doc_id.isalnum():
            return []
        return sorted((self.root / doc_id).glob("page-*.png"))

    def page(self, doc_id: str, index: int) -> Path:
        path = self.root / doc_id / f"page-{index:03}.png"
        if not path.is_file():
            raise FileNotFoundError(path)
        return path

    def _render(self, suffix: str, content: bytes, target: Path) -> list[Path]:
        if suffix == ".pdf":
            pdf = pymupdf.open(stream=content, filetype="pdf")
            if pdf.page_count > 30:
                raise UnsupportedDocument("PDFs are limited to 30 pages")
            pages = []
            for i, page in enumerate(pdf):
                pix = page.get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False)
                path = target / f"page-{i:03}.png"
                pix.save(path)
                pages.append(path)
            return pages
        if suffix in {".txt", ".md", ".docx"}:
            if suffix == ".docx":
                document = Document(io.BytesIO(content))
                text = "\n".join(p.text for p in document.paragraphs)
            else:
                text = content.decode("utf-8", errors="replace")
            return self._render_text(text, target)
        image = Image.open(io.BytesIO(content))
        frames = []
        i = 0
        while True:
            image.seek(i)
            page = ImageOps.exif_transpose(image.convert("RGB"))
            path = target / f"page-{i:03}.png"
            page.save(path, optimize=True)
            frames.append(path)
            i += 1
            try:
                image.seek(i)
            except EOFError:
                break
            if i >= 30:
                break
        return frames

    @staticmethod
    def _render_text(text: str, target: Path) -> list[Path]:
        font = ImageFont.load_default(size=28)
        wrapped: list[str] = []
        for paragraph in text.splitlines() or [""]:
            wrapped.extend(textwrap.wrap(paragraph, width=72) or [""])
        pages = []
        for page_no, start in enumerate(range(0, max(1, len(wrapped)), 46)):
            image = Image.new("RGB", (1240, 1754), "white")
            draw = ImageDraw.Draw(image)
            draw.multiline_text((90, 90), "\n".join(wrapped[start : start + 46]), fill="#17201c", font=font, spacing=12)
            path = target / f"page-{page_no:03}.png"
            image.save(path, optimize=True)
            pages.append(path)
        return pages
